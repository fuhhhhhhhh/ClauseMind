from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path("/home/fuxiangyu/project/ClauseMind")
DELIVERABLE_DIR = ROOT / "doc" / "deliverables"
SCREENSHOT_DIR = ROOT / "doc" / "screenshots"


@dataclass
class ScreenshotSpec:
    filename: str
    title: str
    caption: str


@dataclass
class DocSpec:
    source: Path
    output: Path
    cover_title: str
    heading_title: str
    include_screenshots: bool = False


SCREENSHOTS = [
    ScreenshotSpec("01-login.png", "图 1 登录页面", "用户输入用户名和密码后，可进入 ClauseMind 系统。"),
    ScreenshotSpec("02-register.png", "图 2 注册页面", "新用户可在注册页创建账号，并在成功后进入系统。"),
    ScreenshotSpec("03-dashboard.png", "图 3 工作台页面", "工作台展示系统入口与后端联通状态。"),
    ScreenshotSpec("04-upload.png", "图 4 合同上传页面", "系统支持 PDF、DOCX、TXT 与图片格式合同上传。"),
    ScreenshotSpec("05-contract-list.png", "图 5 合同列表页面", "用户可查看已上传合同，并执行详情、解析或删除操作。"),
    ScreenshotSpec("06-contract-detail.png", "图 6 合同详情页面", "详情页展示合同状态、文件信息以及解析和审查入口。"),
    ScreenshotSpec("07-parse-result.png", "图 7 MinerU 解析结果页面", "系统展示 Markdown 解析结果与解析产物摘要。"),
    ScreenshotSpec("08-normalized-result.png", "图 8 文档标准化结果页面", "标准化结果展示合同标题、主体、章节与条款结构。"),
    ScreenshotSpec("09-review-progress.png", "图 9 审查进度页面", "多 Agent 审查进度页展示流程阶段与执行日志。"),
    ScreenshotSpec("10-risk-analysis.png", "图 10 风险分析页面", "系统按风险等级展示识别出的合同风险项。"),
    ScreenshotSpec("11-suggestion.png", "图 11 修改建议页面", "系统针对风险条款生成对应修改建议与理由。"),
    ScreenshotSpec("12-report.png", "图 12 审查报告页面", "最终报告页展示综合审查结论，并支持导出。"),
    ScreenshotSpec("13-admin.png", "图 13 管理员后台页面", "管理员后台提供统计、用户、合同、任务和 Agent 日志管理功能。"),
]


DOC_SPECS = [
    DocSpec(
        source=DELIVERABLE_DIR / "需求文档.md",
        output=DELIVERABLE_DIR / "ClauseMind系统需求文档.docx",
        cover_title="ClauseMind 系统需求文档",
        heading_title="系统需求文档",
        include_screenshots=True,
    ),
    DocSpec(
        source=DELIVERABLE_DIR / "设计文档.md",
        output=DELIVERABLE_DIR / "ClauseMind系统设计文档.docx",
        cover_title="ClauseMind 系统设计文档",
        heading_title="系统设计文档",
    ),
    DocSpec(
        source=DELIVERABLE_DIR / "项目管理文档.md",
        output=DELIVERABLE_DIR / "ClauseMind项目管理文档.docx",
        cover_title="ClauseMind 项目管理文档",
        heading_title="项目管理文档",
    ),
]


def add_page_number(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def setup_document() -> Document:
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10.5)

    for style_name, font_size in [("Title", 20), ("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        if style_name in styles:
            styles[style_name].font.name = "黑体"
            styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            styles[style_name].font.size = Pt(font_size)

    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    add_page_number(section)
    return doc


def add_cover(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("《面向对象技术与方法》结课设计").bold = True
    p.runs[0].font.size = Pt(20)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(title).bold = True
    p.runs[0].font.size = Pt(24)

    doc.add_paragraph("")
    fields = [
        "学生姓名：付翔宇",
        "项目类型：个人独立开发项目",
        "系统类型：前后端分离的智能化 Web 应用",
        "提交日期：2026 年 6 月 21 日",
    ]
    for line in fields:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(line).font.size = Pt(14)

    doc.add_page_break()


def add_toc_hint(doc: Document) -> None:
    doc.add_heading("目录", level=1)
    doc.add_paragraph("说明：若打开 Word 后目录未自动显示，请在 Word 或 LibreOffice 中选择“更新目录/更新域”。")
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = ' TOC \\o "1-3" \\h \\z \\u '
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    fld_char_text = OxmlElement("w:t")
    fld_char_text.text = "目录将在打开文档后生成"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_text)
    run._r.append(fld_char_end)
    doc.add_page_break()


def normalize_inline_text(text: str) -> str:
    text = text.replace("\\|", "|")
    text = text.replace("&nbsp;", " ")
    text = text.replace("`", "")
    return text.strip()


def add_inline_runs(paragraph, text: str) -> None:
    text = normalize_inline_text(text)
    if not text:
        return

    token_re = re.compile(r"(\*\*.*?\*\*|`.*?`|\*.*?\*)")
    parts = token_re.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`") and len(part) >= 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        elif part.startswith("*") and part.endswith("*") and len(part) >= 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def inline_to_plain_text(text: str) -> str:
    text = normalize_inline_text(text)
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"`(.*?)`", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    return text


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if len(rows) < 2:
        return

    header = rows[0]
    body = [row for row in rows[2:] if len(row) == len(header)]
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    for idx, value in enumerate(header):
        table.rows[0].cells[idx].text = inline_to_plain_text(value)
    for row in body:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = inline_to_plain_text(value)


def add_markdown(doc: Document, text: str) -> None:
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped in {"---", "***"}:
            i += 1
            continue

        if stripped.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i].rstrip())
                i += 1
            p = doc.add_paragraph()
            add_inline_runs(p, "\n".join(code_lines))
            i += 1
            continue

        if re.match(r"^#{1,3}\s+", stripped):
            level = len(stripped.split(" ")[0])
            title = re.sub(r"^#{1,3}\s+", "", stripped)
            p = doc.add_heading("", level=level)
            add_inline_runs(p, title)
            i += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, re.sub(r"^\d+\.\s+", "", stripped))
            i += 1
            continue

        if stripped.startswith(("- ", "* ")):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, stripped[2:])
            i += 1
            continue

        if stripped.startswith(">"):
            p = doc.add_paragraph()
            add_inline_runs(p, stripped.lstrip("> ").strip())
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            rows = []
            for raw in table_lines:
                cells = [c.strip() for c in raw.strip("|").split("|")]
                rows.append(cells)
            add_table(doc, rows)
            continue

        p = doc.add_paragraph()
        add_inline_runs(p, stripped)
        i += 1


def add_screenshot_section(doc: Document) -> None:
    doc.add_heading("2.4 系统运行截图", level=2)
    doc.add_paragraph("以下截图来自本地运行的 ClauseMind 演示环境，使用 demo 与 admin 账号对系统关键功能进行展示。")

    for spec in SCREENSHOTS:
        image_path = SCREENSHOT_DIR / spec.filename
        doc.add_heading(spec.title, level=3)
        if image_path.exists():
            doc.add_picture(str(image_path), width=Cm(15.8))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p = doc.add_paragraph(f"[缺少截图文件：{spec.filename}]")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption = doc.add_paragraph(spec.caption)
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER


def split_requirement_content(markdown: str) -> tuple[str, str]:
    before, marker, after = markdown.partition("## 2.4 系统运行截图")
    if not marker:
        return markdown, ""
    next_section = after.find("## 2.5")
    if next_section == -1:
        return before, ""
    return before, after[next_section:]


def generate_doc(spec: DocSpec) -> None:
    doc = setup_document()
    add_cover(doc, spec.cover_title)
    add_toc_hint(doc)
    doc.add_heading(spec.heading_title, level=1)

    markdown = spec.source.read_text(encoding="utf-8")
    if spec.include_screenshots:
        before, after = split_requirement_content(markdown)
        add_markdown(doc, before)
        add_screenshot_section(doc)
        if after:
            add_markdown(doc, after)
    else:
        add_markdown(doc, markdown)

    doc.save(spec.output)
    print(f"generated {spec.output}")


def main() -> None:
    for spec in DOC_SPECS:
        generate_doc(spec)


if __name__ == "__main__":
    main()
